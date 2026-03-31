"""
JARVIS MAX — RAG Ingestion
Extracts text from PDF, DOCX, TXT, MD, PY, JSON files.
Returns Document objects ready for chunking and embedding.

Usage:
    doc = await ingest_file("/path/to/file.pdf")
    # doc.content  → full extracted text
    # doc.chunks   → list[str] (empty until pipeline chunks them)
"""
from __future__ import annotations

import asyncio
import hashlib
import json
import os
import uuid
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

import structlog

log = structlog.get_logger(__name__)

# Supported extensions → extractor key
_SUPPORTED: dict[str, str] = {
    ".pdf":  "pdf",
    ".docx": "docx",
    ".doc":  "docx",
    ".txt":  "text",
    ".md":   "text",
    ".rst":  "text",
    ".py":   "code",
    ".js":   "code",
    ".ts":   "code",
    ".go":   "code",
    ".rs":   "code",
    ".java": "code",
    ".json": "json",
    ".yaml": "text",
    ".yml":  "text",
    ".toml": "text",
    ".csv":  "text",
}


# ── Data model ────────────────────────────────────────────────

@dataclass
class Document:
    id:       str
    source:   str                     # file path or label
    content:  str                     # full extracted text
    metadata: dict[str, Any]          = field(default_factory=dict)
    chunks:   list[str]               = field(default_factory=list)

    @property
    def content_hash(self) -> str:
        return hashlib.sha256(self.content.encode()).hexdigest()[:16]

    @property
    def word_count(self) -> int:
        return len(self.content.split())

    def to_dict(self) -> dict:
        return {
            "id":         self.id,
            "source":     self.source,
            "content_len": len(self.content),
            "word_count": self.word_count,
            "chunk_count": len(self.chunks),
            "metadata":   self.metadata,
        }


# ── Extractors ────────────────────────────────────────────────

def _extract_pdf(path: str) -> str:
    """pypdf primary, raw binary fallback."""
    try:
        import pypdf  # type: ignore
        reader = pypdf.PdfReader(path)
        pages  = [page.extract_text() or "" for page in reader.pages]
        text   = "\n\n".join(p for p in pages if p.strip())
        log.debug("pdf_extracted_pypdf", pages=len(pages), chars=len(text))
        return text
    except ImportError:
        log.debug("pypdf_unavailable_using_stub")
    except Exception as e:
        log.warning("pypdf_failed", err=str(e)[:80])

    # Raw fallback — extract printable ASCII bytes
    try:
        with open(path, "rb") as f:
            raw = f.read()
        text = "".join(chr(b) for b in raw if 32 <= b < 127 or b in (9, 10, 13))
        log.debug("pdf_extracted_raw", chars=len(text))
        return text
    except Exception as e:
        return f"[PDF extraction failed: {e}]"


def _extract_docx(path: str) -> str:
    """python-docx primary, plain-text fallback."""
    try:
        import docx  # type: ignore  (python-docx)
        doc   = docx.Document(path)
        paras = [p.text for p in doc.paragraphs if p.text.strip()]
        # Include table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paras.append(cell.text.strip())
        text = "\n\n".join(paras)
        log.debug("docx_extracted", paragraphs=len(paras), chars=len(text))
        return text
    except ImportError:
        log.debug("python_docx_unavailable")
    except Exception as e:
        log.warning("docx_failed", err=str(e)[:80])

    # Fallback: try reading as text (will be garbled but better than nothing)
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception as e:
        return f"[DOCX extraction failed: {e}]"


def _extract_json(path: str) -> str:
    """JSON → pretty-printed string."""
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            data = json.load(f)
        return json.dumps(data, indent=2, ensure_ascii=False)
    except Exception as e:
        log.warning("json_extract_failed", path=path, err=str(e)[:80])
        try:
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        except Exception:
            return f"[JSON extraction failed: {e}]"


def _extract_text(path: str) -> str:
    """Plain text / code / markdown."""
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception as e:
        return f"[Text extraction failed: {e}]"


# ── Sync extract ──────────────────────────────────────────────

def _extract_sync(path: str, kind: str) -> str:
    if kind == "pdf":
        return _extract_pdf(path)
    if kind == "docx":
        return _extract_docx(path)
    if kind == "json":
        return _extract_json(path)
    return _extract_text(path)   # text, code


# ── Public API ────────────────────────────────────────────────

async def ingest_file(path: str | Path) -> Document:
    """
    Ingest a file and return a Document.
    Automatically detects type by extension.
    Never raises — on complete failure, content contains error string.
    """
    path_str = str(path)
    ext      = Path(path_str).suffix.lower()
    kind     = _SUPPORTED.get(ext, "text")   # default: read as text

    stat     = None
    try:
        stat = os.stat(path_str)
    except Exception:
        pass

    metadata: dict[str, Any] = {
        "source":    path_str,
        "extension": ext,
        "kind":      kind,
        "size_bytes": stat.st_size if stat else 0,
        "mtime":      stat.st_mtime if stat else 0.0,
    }

    loop    = asyncio.get_running_loop()
    content = await loop.run_in_executor(None, _extract_sync, path_str, kind)

    doc = Document(
        id       = str(uuid.uuid4()),
        source   = path_str,
        content  = content,
        metadata = metadata,
    )
    log.info(
        "document_ingested",
        source=path_str,
        kind=kind,
        chars=len(content),
        words=doc.word_count,
    )
    return doc


async def ingest_text(text: str, source: str = "inline", metadata: dict | None = None) -> Document:
    """Create a Document directly from a text string."""
    return Document(
        id       = str(uuid.uuid4()),
        source   = source,
        content  = text,
        metadata = metadata or {"source": source, "kind": "text"},
    )


def is_supported(path: str | Path) -> bool:
    """Returns True if the file extension is supported."""
    ext = Path(str(path)).suffix.lower()
    return ext in _SUPPORTED
